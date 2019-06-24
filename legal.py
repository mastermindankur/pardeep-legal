import sys
import os
import win32com.client
import xlrd
from docx import Document

def replace_string(doc,old_text,new_text):
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    text = inline[i].text.replace(old_text, new_text)
                    inline[i].text = text
            print (p.text)

    #doc.save('Emplyment-Contract-Filled.docx')
    return doc



document = Document('.\\contract.docx')
wb = xlrd.open_workbook('.\\input_to_contract.xlsx') 
sheet = wb.sheet_by_index(0) 

i=0
for i in range(1,sheet.nrows):
    old_text=str(sheet.cell_value(i, 0)) 
    new_text=str(sheet.cell_value(i, 1))
    print ("Replacing "+ old_text +" with "+new_text)
    #for para in document.paragraphs:
	#print (para.text)
    document=replace_string(document,old_text,new_text)
	
# saving the file now as docx	
document.save('.\\output\\contract-filled.docx')

#saving file as pdf now
wdFormatPDF = 17
in_file = os.path.abspath('.\\output\\contract-filled.docx')
out_file = os.path.abspath('.\\output\\contract-filled.pdf')
word = win32com.client.Dispatch('Word.Application')
doc = word.Documents.Open(in_file)
doc.SaveAs(out_file, FileFormat=wdFormatPDF)
doc.Close()
word.Quit()

